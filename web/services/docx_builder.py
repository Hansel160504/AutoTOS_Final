"""DOCX export — TOS + CILOs + exam items + answer key.

Three-section layout (per panel recommendation):
  Page 1:  TOS (cover, CILO, Table of Specification)
  Page 2+: Exam Items (questions only — NO answers, NO rationale)
  Page N+: Answer Key (compact answer-only list)

Every page also carries the university header (logo + name) and
a 2-column footer (left + right institutional images).

Page: long bond paper 8.5 × 13 in, 1-inch margins, Arial 11pt black-and-white.
"""
from __future__ import annotations

import logging
import os
from io import BytesIO
from typing import List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# Column widths for the TOS table, in twips (1 inch = 1440 twips).
# Column widths for the TOS table, in twips (1 inch = 1440 twips).
# Landscape long bond (13" x 8.5") with 0.7" margins → usable width = 11.6"
_COL_W = [
    int(3.60 * 1440),   # Content Outline
    int(0.90 * 1440),   # Hours
    int(2.20 * 1440),   # FAM
    int(2.20 * 1440),   # INT
    int(2.20 * 1440),   # CRE
    int(0.50 * 1440),   # Total
]  # sum = 16704 twips = 11.6 in# sum = 9360 twips = 6.5 in (matches body width after 1" margins)

# ── Header / footer image paths ─────────────────────────────────────
# These live inside the web container at /app/imgs/.
# Override at runtime via the DOCX_IMG_DIR environment variable.
_DEFAULT_IMG_DIR = "/app/imgs"
IMG_DIR = os.environ.get("DOCX_IMG_DIR", _DEFAULT_IMG_DIR)

HEADER_IMG_PATH    = os.path.join(IMG_DIR, "header_img.png")
FOOTER_LEFT_PATH   = os.path.join(IMG_DIR, "leftmost_img.jpg")
FOOTER_RIGHT_PATH  = os.path.join(IMG_DIR, "rightmost_img.png")

# Image sizes (tweak if header/footer feels too big or small).
HEADER_IMG_HEIGHT  = Inches(0.80)
FOOTER_IMG_HEIGHT  = Inches(0.60)


def build_docx(
    *,
    title: str,
    cilos: List[str],
    topics: List[dict],
    quizzes: List[dict],
    fam_pct: int,
    int_pct: int,
    cre_pct: int,
    total_items: int,
    # NEW (all optional → existing callers/records keep working)
    semester: str = "",          # "1st" | "2nd" | ""
    academic_year: str = "",     # e.g. "2025-2026"
    examination: str = "",       # "Midterm" | "Final" | ""
    descriptive: str = "",       # course name, e.g. "Machine Learning"
    subject: str = "",           # course code,  e.g. "CS328"
    prepared_by: str = "",       # instructor name (uppercased on render)
) -> BytesIO:
    """Public entry point. Returns an in-memory .docx."""
    return _DocxBuilder(
        title=title, cilos=cilos, topics=topics, quizzes=quizzes,
        fam_pct=fam_pct, int_pct=int_pct, cre_pct=cre_pct,
        total_items=total_items,
        semester=semester, academic_year=academic_year,
        examination=examination, descriptive=descriptive,
        subject=subject, prepared_by=prepared_by,
    ).build()


# ──────────────────────────────────────────────────────────────────────
class _DocxBuilder:
    """Owns the Document and writes the three sections in order."""

    _BLACK = RGBColor(0, 0, 0)

    def __init__(
        self, *, title, cilos, topics, quizzes,
        fam_pct, int_pct, cre_pct, total_items,
        semester="", academic_year="", examination="",
        descriptive="", subject="", prepared_by="",
    ):
        self.title = title
        self.cilos = cilos
        self.topics = topics
        self.quizzes = quizzes
        self.fam_pct = fam_pct
        self.int_pct = int_pct
        self.cre_pct = cre_pct
        self.total_items = total_items
        self.semester = semester
        self.academic_year = academic_year
        self.examination = examination
        self.descriptive = descriptive
        self.subject = subject
        self.prepared_by = prepared_by

        self.doc = Document()
        self._configure_page()
        self._configure_default_style()
        self._add_header_and_footer()

    # ── Public ─────────────────────────────────────────────────
    # ── Public ─────────────────────────────────────────────────
    def build(self) -> BytesIO:
        # ── Page 1: TOS (title, meta, CILO, table, signatures, note) ──
        self._add_title_header()
        self._add_metadata_section()
        if self.cilos:
            self._add_cilos_section()
        self._add_tos_table()
        self._add_signatures_section()
        self._add_form_note()

        # ── Page 2+: Exam Items (UNCHANGED) ──
        self._add_page_break()
        self._add_exam_items(include_answers=False)

        # ── Page N+: Answer Key (UNCHANGED) ──
        self._add_page_break()
        self._add_answer_key()

        buf = BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf
    # ── Page / style config ────────────────────────────────────
    def _configure_page(self) -> None:
        sec = self.doc.sections[0]
        sec.orientation = WD_ORIENTATION.LANDSCAPE
        # Long bond paper, landscape: 13" wide × 8.5" tall
        sec.page_width = Inches(13)
        sec.page_height = Inches(8.5)
        # Tighter margins to fit metadata + table + signatures on page 1
        sec.top_margin = sec.bottom_margin = Inches(0.6)
        sec.left_margin = sec.right_margin = Inches(0.7)
        sec.header_distance = Inches(0.3)
        sec.footer_distance = Inches(0.3)

    def _configure_default_style(self) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(11)
        normal.font.color.rgb = self._BLACK

    # ── Header & footer (appears on every page) ────────────────
    def _add_header_and_footer(self) -> None:
        section = self.doc.sections[0]
        self._build_header(section)
        self._build_footer(section)

    def _build_header(self, section) -> None:
        """Centered: logo image + 'Republic of the Philippines' +
        'North Eastern Mindanao State University'."""
        header = section.header

        # Clear default empty paragraph
        if header.paragraphs:
            header.paragraphs[0].clear()
            p_logo = header.paragraphs[0]
        else:
            p_logo = header.add_paragraph()

        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after = Pt(2)

        # Image (skip silently if missing)
        if os.path.isfile(HEADER_IMG_PATH):
            try:
                run = p_logo.add_run()
                run.add_picture(HEADER_IMG_PATH, height=HEADER_IMG_HEIGHT)
            except Exception as exc:
                logger.warning("Header image failed to load: %s", exc)
        else:
            logger.info("Header image not found at %s — skipping", HEADER_IMG_PATH)

        # "Republic of the Philippines"
        p_country = header.add_paragraph()
        p_country.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_country.paragraph_format.space_before = Pt(0)
        p_country.paragraph_format.space_after = Pt(0)
        self._run(p_country, "Republic of the Philippines",
                  bold=True, size=11)

        # "North Eastern Mindanao State University"
        p_univ = header.add_paragraph()
        p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_univ.paragraph_format.space_before = Pt(0)
        p_univ.paragraph_format.space_after = Pt(0)
        self._run(p_univ, "North Eastern Mindanao State University",
                  bold=True, size=12)

    def _build_footer(self, section) -> None:
        """2-column borderless table — leftmost image on the left,
        rightmost image on the right."""
        footer = section.footer

        # Clear default paragraph
        if footer.paragraphs:
            footer.paragraphs[0].clear()

        # Two equal-width columns spanning the full body width (6.5 in)
        tbl = footer.add_table(rows=1, cols=2, width=Inches(6.5))
        tbl.autofit = False
        self._set_table_no_borders(tbl)

        # Set cell widths in twips (1 inch = 1440 twips, so 3.25 in = 4680)
        half_twips = int(3.25 * 1440)

        left_cell, right_cell = tbl.rows[0].cells
        for cell in (left_cell, right_cell):
            self._set_cell_width(cell, half_twips)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].text = ""

        # Left image — left-aligned
        if os.path.isfile(FOOTER_LEFT_PATH):
            p_left = left_cell.paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                p_left.add_run().add_picture(FOOTER_LEFT_PATH,
                                             height=FOOTER_IMG_HEIGHT)
            except Exception as exc:
                logger.warning("Footer left image failed to load: %s", exc)
        else:
            logger.info("Footer left image not found at %s — skipping",
                        FOOTER_LEFT_PATH)

        # Right image — right-aligned
        if os.path.isfile(FOOTER_RIGHT_PATH):
            p_right = right_cell.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                p_right.add_run().add_picture(FOOTER_RIGHT_PATH,
                                              height=FOOTER_IMG_HEIGHT)
            except Exception as exc:
                logger.warning("Footer right image failed to load: %s", exc)
        else:
            logger.info("Footer right image not found at %s — skipping",
                        FOOTER_RIGHT_PATH)


    @staticmethod
    def _set_table_no_borders(table) -> None:
        """Remove all borders from a table (used for the footer layout)."""
        tbl_pr = table._tbl.tblPr
        tbl_borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right",
                     "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            tbl_borders.append(b)
        tbl_pr.append(tbl_borders)

    # ── Low-level helpers ──────────────────────────────────────
    def _run(self, para, text: str, *, bold=False, size=11, italic=False):
        r = para.add_run(text or "")
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = self._BLACK
        return r

    def _add_para(
        self,
        text: str = "",
        *,
        bold: bool = False,
        size: int = 11,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        italic: bool = False,
        space_before: int = 0,
        space_after: int = 4,
    ):
        p = self.doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            self._run(p, text, bold=bold, size=size, italic=italic)
        return p

    def _add_page_break(self) -> None:
        """Insert a hard page break."""
        p = self.doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    def _cell_write(
        self, cell, text: str,
        *, bold=False, size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=False,
        valign=WD_ALIGN_VERTICAL.CENTER,
    ) -> None:
        cell.vertical_alignment = valign
        lines = (text or "").split("\n")
        p0 = cell.paragraphs[0]
        p0.alignment = align
        p0.paragraph_format.space_before = Pt(1)
        p0.paragraph_format.space_after = Pt(1)
        p0.paragraph_format.left_indent = Inches(0.05)
        p0.paragraph_format.right_indent = Inches(0.05)
        self._run(p0, lines[0], bold=bold, size=size, italic=italic)
        for line in lines[1:]:
            p = cell.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.05)
            p.paragraph_format.right_indent = Inches(0.05)
            self._run(p, line, bold=bold, size=size, italic=italic)

    @staticmethod
    def _shade_cell(cell, hex_color: str = "E0E0E0") -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    @staticmethod
    def _set_cell_width(cell, twips: int) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = OxmlElement("w:tcW")
        tc_w.set(qn("w:w"), str(twips))
        tc_w.set(qn("w:type"), "dxa")
        tc_pr.append(tc_w)

    @staticmethod
    def _set_table_width(table, twips: int) -> None:
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), str(twips))
        tbl_w.set(qn("w:type"), "dxa")
        table._tbl.tblPr.append(tbl_w)

    @staticmethod
    def _underline_paragraph(p) -> None:
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        borders.append(bottom)
        p_pr.append(borders)

    # ── Title header at the top of page 1 ──────────────────────
    # ── Title header at the top of page 1 ──────────────────────
    def _add_title_header(self) -> None:
        """Document title above the metadata strip."""
        self._add_para(
            "TABLE OF SPECIFICATIONS",
            bold=True, size=14,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=0, space_after=6,
        )
        if self.title:
            self._add_para(
                self.title,
                bold=True, size=11, italic=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=6,
            )
    # ── Section 1: CILOs ───────────────────────────────────────
    # ── Section 1: CILOs ───────────────────────────────────────
    def _add_cilos_section(self) -> None:
        self._add_para(
            "COGNITIVE OBJECTIVES / BEHAVIORAL DIMENSIONS / THINKING SKILLS",
            bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=4, space_after=2,
        )
        self._add_para(
            "Course Intended Learning Outcomes (CILO):",
            bold=True, size=10, space_before=2, space_after=3,
        )
        for i, c in enumerate(self.cilos, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            self._run(p, f"CLO{i}.  {c}", size=10)
        self._add_para(space_after=6)
    # ── Section 2: TOS Table ───────────────────────────────────
    def _add_tos_table(self) -> None:
        self._add_para(
            "TABLE OF SPECIFICATION",
            bold=True, size=11,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=6,
        )

        n = len(self.topics)
        tbl = self.doc.add_table(rows=3 + n + 1, cols=6)
        tbl.style = "Table Grid"
        self._set_table_width(tbl, sum(_COL_W))

        # Set column widths on every cell (required for vertical-merge accuracy).
        for row in tbl.rows:
            for ci, cell in enumerate(row.cells):
                self._set_cell_width(cell, _COL_W[ci])

        # Merge "Content Outline" & "Hours" across the three header rows.
        tbl.cell(0, 0).merge(tbl.cell(2, 0))
        tbl.cell(0, 1).merge(tbl.cell(2, 1))

        self._write_tos_header(tbl)
        self._write_tos_data(tbl, start_row=3)
        self._add_para(space_after=14)

    def _write_tos_header(self, tbl) -> None:
        # Row 0 — column titles.
        self._cell_write(tbl.cell(0, 0), "CONTENT OUTLINE",
                         bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._cell_write(tbl.cell(0, 1), "NUMBER OF\nHOURS SPENT",
                         bold=True, size=9)
        self._cell_write(tbl.cell(0, 2),
                         "FAMILIARIZATION\n(Remembering / Understanding)",
                         bold=True, size=9)
        self._cell_write(tbl.cell(0, 3),
                         "INTEGRATION\n(Applying / Analyzing)",
                         bold=True, size=9)
        self._cell_write(tbl.cell(0, 4),
                         "CREATION\n(Evaluating / Creating)",
                         bold=True, size=9)
        self._cell_write(tbl.cell(0, 5), "TOTAL", bold=True, size=9)
        for ci in range(6):
            self._shade_cell(tbl.cell(0, ci), "D0D0D0")

        # Row 1 — percentages.
        self._cell_write(tbl.cell(1, 2), f"(Percentage)\n{self.fam_pct}%",
                         size=9, italic=True)
        self._cell_write(tbl.cell(1, 3), f"(Percentage)\n{self.int_pct}%",
                         size=9, italic=True)
        self._cell_write(tbl.cell(1, 4), f"(Percentage)\n{self.cre_pct}%",
                         size=9, italic=True)
        self._cell_write(tbl.cell(1, 5), "100%", size=9, italic=True)
        for ci in range(2, 6):
            self._shade_cell(tbl.cell(1, ci), "EBEBEB")

        # Row 2 — sub-header labels.
        for col, label in ((2, "Item Numbers"), (3, "Item Numbers"),
                           (4, "Item Numbers"), (5, "Total No.\nof Items")):
            self._cell_write(tbl.cell(2, col), label, bold=True, size=9)
        for ci in range(2, 6):
            self._shade_cell(tbl.cell(2, ci), "EBEBEB")

    def _write_tos_data(self, tbl, *, start_row: int) -> None:
        t_hrs = t_fam = t_int = t_cre = t_tot = 0
        for i, t in enumerate(self.topics):
            row = tbl.rows[start_row + i]
            hrs = t.get("hours") or 0
            fam = t.get("fam") or 0
            intg = t.get("int") or 0
            cre = t.get("cre") or 0
            tot = t.get("items") or t.get("quiz_items") or 0
            t_hrs += hrs; t_fam += fam; t_int += intg
            t_cre += cre; t_tot += tot

            self._cell_write(row.cells[0], t.get("topic") or "",
                             bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
            self._cell_write(row.cells[1], str(hrs), size=10)
            self._cell_write(row.cells[2], t.get("fam_range") or "—", size=10)
            self._cell_write(row.cells[3], t.get("int_range") or "—", size=10)
            self._cell_write(row.cells[4], t.get("cre_range") or "—", size=10)
            self._cell_write(row.cells[5], str(tot), bold=True, size=10)

        # Totals row.
        footer = tbl.rows[start_row + len(self.topics)]
        self._cell_write(footer.cells[0], "TOTAL:",
                         bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
        self._cell_write(footer.cells[1], str(t_hrs), bold=True, size=10)
        self._cell_write(footer.cells[2], str(t_fam), bold=True, size=10)
        self._cell_write(footer.cells[3], str(t_int), bold=True, size=10)
        self._cell_write(footer.cells[4], str(t_cre), bold=True, size=10)
        self._cell_write(footer.cells[5], str(t_tot), bold=True, size=10)
        for ci in range(6):
            self._shade_cell(footer.cells[ci], "D0D0D0")

    # ── Section 3: Exam Items ──────────────────────────────────
    def _add_exam_items(self, include_answers: bool = True) -> None:
        """Render the exam items. When include_answers=False, the answer line
        and rationale are omitted (used for the student-facing exam page)."""
        self._add_para(
            "EXAM ITEMS",
            bold=True, size=14,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=0, space_after=4,
        )
        if not include_answers:
            self._add_para(
                "Name: _______________________________     "
                "Score: ________     Date: _______________",
                size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=10,
            )

        current_test: Optional[str] = ""
        for idx, q in enumerate(self.quizzes, 1):
            if not isinstance(q, dict):
                continue

            header = q.get("test_header") or ""
            desc = q.get("test_description") or ""

            if header and header != current_test:
                current_test = header
                self._write_test_header(header, desc)

            self._write_question(idx, q, include_answer=include_answers)

    def _write_test_header(self, header: str, desc: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        self._run(p, header, bold=True, size=12)
        if desc:
            self._run(p, f"   —   {desc}", italic=True, size=11)
        self._underline_paragraph(p)

    def _write_question(self, idx: int, q: dict, *, include_answer: bool = True) -> None:
        qtype = (q.get("type") or "").lower()

        p_q = self.doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(4)
        p_q.paragraph_format.space_after = Pt(2)
        p_q.paragraph_format.keep_with_next = True
        self._run(p_q, f"{idx}.  {q.get('question') or ''}", size=11)

        if qtype == "mcq" and q.get("choices"):
            self._write_mcq_choices(q, highlight_answer=include_answer)
        elif qtype in ("truefalse", "tf", "true_false"):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            self._run(p, "(True or False)", italic=True, size=11)
        elif qtype in ("open_ended", "open-ended", "open ended"):
            if not include_answer:
                for _ in range(3):
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.35)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    self._run(p, "_" * 90, size=11)

        if include_answer:
            self._write_answer_line(q)

    def _write_mcq_choices(self, q: dict, *, highlight_answer: bool = True) -> None:
        letters = ["a", "b", "c", "d", "e"]
        ans_raw = (q.get("answer") or "").strip()
        ans_low = ans_raw.lower()
        ans_upper = ans_raw.upper().rstrip(".")
        ans_is_letter = ans_upper in ("A", "B", "C", "D")

        for ci, choice in enumerate(q["choices"][:5]):
            letter = letters[ci]
            choice_l = (choice or "").lower()
            is_correct = False
            if highlight_answer:
                if ans_is_letter:
                    is_correct = (ans_upper == letter.upper())
                else:
                    is_correct = (
                        choice_l == ans_low
                        or (ans_low and len(ans_low) > 5 and ans_low in choice_l)
                        or ans_low == letter
                    )

            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            self._run(p, f"{letter})  {choice or ''}", bold=is_correct, size=11)

    def _write_answer_line(self, q: dict) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(7)
        self._run(p, "Answer: ", bold=True, size=11)

        qtype = (q.get("type") or "").lower()
        ans_raw = str(q.get("answer") or "").strip()
        ans_upper = ans_raw.upper().rstrip(".")

        if qtype == "mcq" and ans_upper in ("A", "B", "C", "D"):
            choices = q.get("choices") or []
            idx = ord(ans_upper) - ord("A")
            if 0 <= idx < len(choices):
                self._run(p, f"{ans_upper}) {choices[idx]}", size=11)
            else:
                self._run(p, ans_raw or "—", size=11)
        else:
            self._run(p, ans_raw or "—", size=11)

        if q.get("answer_text") and not q.get("_invalid_tf"):
            self._run(p, f"  ({q['answer_text']})", italic=True, size=10)

    # ── Section 4: Answer Key ──────────────────────────────────
    def _add_answer_key(self) -> None:
        self._add_para(
            "ANSWER KEY",
            bold=True, size=14,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=0, space_after=2,
        )
        self._add_para(
            "FOR INSTRUCTOR USE ONLY",
            italic=True, size=9,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=10,
        )

        current_test: Optional[str] = ""
        for idx, q in enumerate(self.quizzes, 1):
            if not isinstance(q, dict):
                continue

            header = q.get("test_header") or ""
            if header and header != current_test:
                current_test = header
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.keep_with_next = True
                self._run(p, header, bold=True, size=11)
                self._underline_paragraph(p)

            self._write_answer_key_line(idx, q)

    def _write_answer_key_line(self, idx: int, q: dict) -> None:
        qtype = (q.get("type") or "").lower()
        ans_raw = str(q.get("answer") or "").strip()
        ans_upper = ans_raw.upper().rstrip(".")

        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.10)
        p.paragraph_format.first_line_indent = Inches(-0.10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)

        self._run(p, f"{idx}.  ", bold=True, size=11)

        if qtype == "mcq" and ans_upper in ("A", "B", "C", "D"):
            choices = q.get("choices") or []
            choice_idx = ord(ans_upper) - ord("A")
            if 0 <= choice_idx < len(choices):
                self._run(p, f"{ans_upper}) {choices[choice_idx]}", size=11)
            else:
                self._run(p, ans_raw or "—", size=11)
        elif qtype in ("truefalse", "tf", "true_false"):
            display = ans_upper if ans_upper in ("TRUE", "FALSE") else ans_raw
            self._run(p, display or "—", size=11)
        elif qtype in ("open_ended", "open-ended", "open ended"):
            self._run(p, ans_raw or "—", size=10)
        else:
            self._run(p, ans_raw or "—", size=11)

        if q.get("answer_text") and not q.get("_invalid_tf") and qtype != "open_ended":
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.35)
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(3)
            self._run(p2, q["answer_text"], italic=True, size=9)

    # ── Section 1.5: Metadata strip (semester / exam / subject) ─
    def _add_metadata_section(self) -> None:
        """Two stacked 2-column borderless rows below the title."""
        check = lambda v: "\u2611" if v else "\u2610"  # ☑ / ☐
        ay = self.academic_year or "________"
        sem1 = check(self.semester.lower().startswith("1"))
        sem2 = check(self.semester.lower().startswith("2"))
        ex_mid = check(self.examination.lower().startswith("mid"))
        ex_fin = check(self.examination.lower().startswith("fin"))

        tbl = self.doc.add_table(rows=2, cols=2)
        tbl.autofit = False
        self._set_table_no_borders(tbl)
        self._set_table_width(tbl, sum(_COL_W))

        half = sum(_COL_W) // 2
        for r in tbl.rows:
            for cell in r.cells:
                self._set_cell_width(cell, half)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Row 1: Semester/AY  |  Examination
        p = tbl.rows[0].cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        self._run(p, "Semester, Academic Year: ", bold=True, size=10)
        self._run(p, f"  {sem1} 1st Sem AY {ay}    {sem2} 2nd Sem AY {ay}", size=10)

        p = tbl.rows[0].cells[1].paragraphs[0]
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        self._run(p, "Examination: ", bold=True, size=10)
        self._run(p, f"  {ex_mid} Midterm    {ex_fin} Final", size=10)

        # Row 2: Descriptive | Subject
        p = tbl.rows[1].cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(0)
        self._run(p, "Descriptive: ", bold=True, size=10)
        self._run(p, self.descriptive or "________", size=10)

        p = tbl.rows[1].cells[1].paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(0)
        self._run(p, "Subject: ", bold=True, size=10)
        self._run(p, self.subject or "________", size=10)

        self._add_para(space_after=4)

    # ── Section 2.5: Signatures (Prepared / Reviewed / Approved) ─
    # Reviewed by + Approved by are FIXED per institutional policy.
    _REVIEWED_NAME = "CHRISTINE W. PITOS, MSCS"
    _APPROVED_NAME = "LOLITA M. MARTIN, PhD(CAR)"

    def _add_signatures_section(self) -> None:
        prepared = (self.prepared_by or "").strip().upper() or "________________________"

        # — Prepared by —
        self._add_para("Prepared by:", bold=True, size=10,
                       space_before=10, space_after=18)
        self._add_para(prepared, bold=True, size=10,
                       align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
        self._add_para("Instructor", size=10, italic=True, space_after=10)

        # — Reviewed by (two signatories, side by side) —
        self._add_para("Reviewed by:", bold=True, size=10,
                       space_before=4, space_after=18)
        tbl = self.doc.add_table(rows=2, cols=2)
        tbl.autofit = False
        self._set_table_no_borders(tbl)
        self._set_table_width(tbl, sum(_COL_W))
        half = sum(_COL_W) // 2
        for r in tbl.rows:
            for cell in r.cells:
                self._set_cell_width(cell, half)

        self._cell_write(tbl.cell(0, 0), self._REVIEWED_NAME,
                         bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        self._cell_write(tbl.cell(0, 1), self._REVIEWED_NAME,
                         bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        self._cell_write(tbl.cell(1, 0), "Program Coordinator",
                         italic=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        self._cell_write(tbl.cell(1, 1), "Department Chair",
                         italic=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        self._add_para(space_after=8)

        # — Approved by —
        self._add_para("Approved by:", bold=True, size=10,
                       space_before=4, space_after=18)
        self._add_para(self._APPROVED_NAME, bold=True, size=10,
                       align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
        self._add_para("Assistant Campus Director",
                       italic=True, size=10, space_after=6)

    # ── Section 2.7: Form-code note ─────────────────────────────
    def _add_form_note(self) -> None:
        self._add_para(
            "Note: Allocation of the percentage will be identified by the "
            "faculty based on what is outlined on the syllabus.",
            italic=True, size=9, space_before=4, space_after=2,
        )
        self._add_para(
            "FM-ACAD-007/REV003/02.05.2025",
            italic=True, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT,
            space_after=0,
        )
